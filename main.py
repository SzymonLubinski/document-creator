import pandas as pd
import numpy as np
import matplotlib.pyplot as plt
from matplotlib import transforms
import seaborn as sns
# from docx import Document
# from docx.shared import Inches
import docx
from docx2pdf import convert


def data_preparation(dataframe):
    # Columns names changed
    provinces = dataframe['Nazwa'].tolist()
    dataframe = dataframe.transpose()
    dataframe.drop(['Nazwa'], inplace=True)
    for i in range(len(provinces)):
        dataframe.rename(columns={i: provinces[i]}, inplace=True)

    # Index reconstruction
    dataframe.reset_index(inplace=True)
    separated = dataframe['index'].str.split(';', expand=True)
    new_cols_names = ['quarter', 'to_remove', 'area', 'year', 'to_remove']
    for i in range(len(separated.columns)):
        separated.rename(columns={i: new_cols_names[i]}, inplace=True)

    # Set multiindex
    dataframe = pd.concat([dataframe, separated], axis=1, join='inner')
    dataframe.rename(columns={'index': 'to_remove'}, inplace=True)
    dataframe.drop(columns=['to_remove'], inplace=True)
    dataframe.set_index(['year', 'quarter', 'area'], inplace=True)
    dataframe.sort_index(inplace=True)

    return dataframe


def set_area(a):
    r = average_price.swaplevel(0, 2)
    # r = r.swaplevel(1, 2)
    r = r[region].loc[a]
    r = r.reset_index()
    r.rename(columns={region: 'average_price'}, inplace=True)
    # print('set')
    # print(r.head(7))
    return r


def all_area():
    r = average_price[region]
    r = r.reset_index()
    r = r.pivot_table(columns=['quarter', 'year'], values=region, aggfunc='mean')
    r = r.transpose()
    r = r.reset_index()
    r.rename(columns={region: 'average_price'}, inplace=True)
    # print('all')
    # print(r.head(7))
    return r


# File read
average_price = pd.read_csv('1m2_average.csv', sep=';')
average_price = average_price.drop(columns=['Kod', 'Unnamed: 82'])

# Second dataframe
apartments_sold = pd.read_csv('apartments_sold.csv', sep=';')
apartments_sold = apartments_sold.drop(columns=['Kod', 'Unnamed: 22'])

# Calling
average_price = data_preparation(average_price)
apartments_sold = data_preparation(apartments_sold)
apartments_sold = apartments_sold.reset_index(level=2, drop=True)

# PODKARPACKIE in quarter
region = input(f'Choose one of regions: {average_price.columns.tolist()}')
if region not in average_price.columns.tolist():
    region = 'PODKARPACKIE'

# Size of the apartments
chosen_area = input('Choose area in m2 unit: -40, 40-60, 60-80 80+ or all')
match chosen_area:
    case '-40':
        selected_region = set_area('do 40 m2')
    case '40-60':
        selected_region = set_area('od 40,1 do 60 m2')
    case '60-80':
        selected_region = set_area('od 60,1 do 80 m2')
    case '80+':
        selected_region = set_area('od 80,1 m2')
    case 'all':
        selected_region = all_area()
    case _:
        selected_region = all_area()


# PODKARPACKIE apartments sold
apartments_region = apartments_sold[region]
apartments_region = apartments_region.reset_index()
apartments_region.rename(columns={region: 'apartments_sold'}, inplace=True)


def losing_weight(fat):
    fat['quarter'] = fat['quarter'].astype('category')
    fat['year'] = fat['year'].astype('category')
    fat[fat.columns[2]] = fat[fat.columns[2]].astype('int64')


losing_weight(selected_region)
losing_weight(apartments_region)

# Matplotlib - first chart
quarters = []
for i in pd.unique(selected_region['quarter']):
    q = 'q' + i[0]
    quarters.append(q)
years = pd.unique(selected_region['year'])

fig, ax = plt.subplots(figsize=(12, 4))
ax = sns.barplot(x='year', y='average_price', hue='quarter', palette='Reds', data=selected_region, ax=ax)
year_pos = np.sort([p.get_x() + p.get_width() / 2 for p in ax.patches])
ax.set_xticks(year_pos)
ax.set_xticklabels(np.tile(quarters, len(years)), rotation=30)
ax.get_legend().remove()
ax.set_xlabel('')
fruit_pos = year_pos.reshape(-1, len(quarters)).mean(axis=1)
trans = transforms.blended_transform_factory(ax.transData, ax.transAxes)
for pos, label in zip(fruit_pos, years):
    ax.text(pos, -0.25, label, transform=trans, ha='center', va='bottom', color='steelblue', fontsize=14)
for pos in (fruit_pos[:-1] + fruit_pos[1:]) / 2:
    ax.axvline(pos, 0, -0.25, color='steelblue', ls=':', clip_on=False)
ax.spines['top'].set_visible(False)
ax.spines['right'].set_visible(False)
plt.tight_layout()
plt.savefig('foo.png')
plt.clf()

# Create document
start_period = years[0]
end_period = years[-1]
report = docx.Document()
doc_styles = report.styles['Normal']
font = doc_styles.font
font.name = 'Arial'
font.size = docx.shared.Pt(12)


# Head of first page
head = f'Housing prices in the {region.title()}'
report.add_heading(head, 0)

# Paragraph of document
p = report.add_paragraph(
    f'This report determines the average price per square meter of flats in the {region.title()} region. '
    'The time period covers the years from ')
p.add_run(start_period).bold = True
p.add_run(' to ')
p.add_run(end_period).bold = True
p.add_run(' broken down into quarters.')

# Add picture
report.add_picture('foo.png', width=docx.shared.Inches(6.5))

# Statistic data
# The biggest price
min_price = selected_region.loc[selected_region['average_price'].astype(float).idxmin()]
max_price = selected_region.loc[selected_region['average_price'].astype(float).idxmax()]

p = report.add_paragraph(f'The highest price was PLN {max_price["average_price"]} in Q{max_price["quarter"][0]} '
                         f'{max_price["year"]}.')
last_quarter = pd.unique(selected_region['quarter'])[-1][0]
if max_price["quarter"][0] == last_quarter and max_price["year"] == end_period:
    text_value = ' This was the last period studied.'
else:
    text_value = ' In subsequent periods, average prices were lower.'
p.add_run(text_value)
p.add_run(f' The lowest price was PLN {min_price["average_price"]} in Q{min_price["quarter"][0]} {min_price["year"]}.')

# The difference between the first and last period
first = selected_region.iloc[0]
last = selected_region.iloc[-1]
difference = last["average_price"] - first["average_price"]
p = report.add_paragraph(f'In the first period under review, the price was PLN {first["average_price"]}, '
                         f'while in the last period the price was PLN {last["average_price"]}.')
if difference > 0:
    text_value = f' The price in the analyzed period increased by PLN {difference}.'
else:
    text_value = f' The price in the analyzed period fell by PLN {abs(difference)}.'
p.add_run(text_value)

# Differences between adjacent periods
previous_value = None
max_difference = 0
max_difference_index = 'none'

for index in selected_region.index:
    if previous_value is not None:
        if abs(selected_region["average_price"][index] - previous_value) > max_difference:
            max_difference = abs(selected_region["average_price"][index] - previous_value)
            max_difference_index = index
    else:
        previous_value = selected_region["average_price"][index]
    previous_value = selected_region["average_price"][index]

before_diff = selected_region.iloc[max_difference_index - 1]
after_diff = selected_region.iloc[max_difference_index]
p = report.add_paragraph(f'The biggest change in price between the quarters under study occurred between '
                         f'Q{before_diff["quarter"][0]} 'f'{before_diff["year"]} and Q{after_diff["quarter"][0]} '
                         f'{after_diff["year"]}.')

if after_diff["average_price"] - before_diff["average_price"] > 0:
    text_value = f' The value in this period increased by PLN {max_difference}.'
else:
    text_value = f' T he value during this period decreased by PLN {max_difference}.'
p.add_run(text_value)

# Second page
report.add_page_break()


# Data comparison by index 2017 Q1 = 100%
def indexing100(ind, value):
    x = 100 * value / ind
    return x


dependence = apartments_region.merge(selected_region, how='right')
dependence.insert(0, 'period', [dependence.loc[index]['year'] + ' Q' + dependence.loc[index]['quarter'][0] for index in
                                dependence.index])
dependence.set_index('period', inplace=True)

price_index = dependence['average_price'].iloc[0]
apartments_index = dependence['apartments_sold'].iloc[0]
dependence['price_shift'] = dependence.apply(lambda x: indexing100(price_index, x['average_price']), axis=1)
dependence['apartments_shift'] = dependence.apply(lambda x: indexing100(apartments_index, x['apartments_sold']), axis=1)

# Matplotlib - second chart
plt.plot(dependence['price_shift'], label='Cost per m2')
plt.plot(dependence['apartments_shift'], label='apartments sold')
plt.legend(loc='upper left')
plt.ylabel('in percents %')
plt.xticks(rotation=45)
plt.tight_layout()
plt.savefig('foo2.png')
# plt.show()

# Head of second page
head = 'Average price and number of sold apartments'
report.add_heading(head, 0)

p = report.add_paragraph('The chart shows the ratio of the average price of flats per 1 m2 to the number of flats '
                         'sold. As the starting index, the first analyzed period, 2017 Q1 = 100%.')
report.add_picture('foo2.png', width=docx.shared.Inches(6.5))

# Counting shifts
common_loss = 0
common_increase = 0
reverse_action = 0
action_dict = {
    'common_increase': 0,
    'common_loss': 0,
    'reverse_action': 0
}
last_index = {0: dependence.iloc[0]['price_shift'], 1: dependence.iloc[0]['apartments_shift']}

for index in dependence.index[1:]:
    if dependence.loc[index]['price_shift'] > last_index[0] and dependence.loc[index]['apartments_shift'] > last_index[1]:
        action_dict['common_increase'] += 1
    elif dependence.loc[index]['price_shift'] < last_index[0] and dependence.loc[index]['apartments_shift'] < last_index[1]:
        action_dict['common_loss'] += 1
    else:
        action_dict['reverse_action'] += 1

    last_index[0] = dependence.loc[index]['price_shift']
    last_index[1] = dependence.loc[index]['apartments_shift']

if max(action_dict) == 'common_increase':
    text_value = 'the average price of flats per 1 m2  usually' \
                 'increased along with the increase in the number of flats sold.'
if max(action_dict) == 'common_loss':
    text_value = 'the average price of flats per 1 m2  usually ' \
                 'decreased along with the decrease in the number of flats sold.'
if max(action_dict) == 'reverse_action':
    text_value = 'the decrease or increase in the sale of flats most often had ' \
                 'the opposite effect on the price of flats per 1 m2.'

p = report.add_paragraph(f'In the analyzed period, {text_value}')

# Save document
report.save('price report.docx')
# convert('price report.docx', 'price report.pdf')
